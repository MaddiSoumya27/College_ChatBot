"""
ingest.py
---------
Phase 1: Load BVRITH_Knowledge_Base.docx → chunk → embed → persist to Chroma.

Primary source: BVRITH_Knowledge_Base.docx  (always reflects your latest edits)
Fallback:       kb/*.md files               (used only if .docx is missing)

Run once:  python3 ingest.py
Re-check:  python3 ingest.py --check
"""

import argparse
import os
import re
import shutil
from datetime import date

from dotenv import load_dotenv
load_dotenv()

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings

# ── Config ────────────────────────────────────────────────────────────────────
BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
CHROMA_DIR  = os.path.join(BASE_DIR, "bvrith_chroma_db")
DOCX_PATH   = os.path.join(BASE_DIR, "BVRITH_Knowledge_Base.docx")
KB_DIR      = os.path.join(BASE_DIR, "kb")

CHUNK_SIZE    = 500
CHUNK_OVERLAP = 75
EMBED_MODEL   = os.getenv("EMBED_MODEL", "text-embedding-3-small")
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", None)
LAST_VERIFIED = str(date.today())

# Canonical H1 section names (must match headings in the .docx exactly)
SECTION_ORDER = [
    "About BVRIT",
    "Departments",
    "Admissions",
    "Fee Structure",
    "Placements",
    "Campus & Facilities",
    "Faculty",
    "Contact",
    "Live Site Data",
]

# ── DOCX loader ───────────────────────────────────────────────────────────────

def load_docx_documents() -> list[Document]:
    """
    Parse BVRITH_Knowledge_Base.docx section-by-section.
    Each Heading-1 block becomes one LangChain Document with section metadata.
    Heading-2 and Heading-3 text is preserved inline so the splitter can use
    them as natural split points.
    """
    from docx import Document as DocxDocument  # python-docx

    print(f"  Reading {DOCX_PATH} …")
    doc = DocxDocument(DOCX_PATH)

    # Walk paragraphs, split on H1 boundaries
    sections: dict[str, list[str]] = {}
    current_h1 = "General"

    for para in doc.paragraphs:
        style = para.style.name
        text  = para.text.strip()

        if not text:
            continue

        if style == "Heading 1":
            current_h1 = text
            sections.setdefault(current_h1, [])
            sections[current_h1].append(f"# {text}")
        elif style == "Heading 2":
            sections.setdefault(current_h1, [])
            sections[current_h1].append(f"\n## {text}")
        elif style == "Heading 3":
            sections.setdefault(current_h1, [])
            sections[current_h1].append(f"\n### {text}")
        else:
            sections.setdefault(current_h1, [])
            sections[current_h1].append(text)

    # Build Documents in canonical order (unknown sections appended at the end)
    docs = []
    ordered_keys = [k for k in SECTION_ORDER if k in sections]
    extra_keys   = [k for k in sections if k not in SECTION_ORDER and k != "General"]

    for key in ordered_keys + extra_keys:
        content = "\n".join(sections[key]).strip()
        if not content:
            continue
        docs.append(Document(
            page_content=content,
            metadata={"section_hint": key, "source": "BVRITH_Knowledge_Base.docx"}
        ))
        print(f"    Loaded section '{key}' — {len(content):,} chars")

    print(f"  Loaded {len(docs)} sections from docx")
    return docs


# ── Markdown fallback loader ───────────────────────────────────────────────────
KB_FILES = [
    ("About BVRIT",         "00_college_overview.md"),
    ("Departments",         "02_departments.md"),
    ("Admissions",          "01_admissions.md"),
    ("Fee Structure",       "01_admissions.md"),
    ("Placements",          "03_placements.md"),
    ("Campus & Facilities", "04_facilities_research_governance.md"),
    ("Faculty",             "05_faculty_directory.md"),
    ("Contact",             "04_facilities_research_governance.md"),
    ("Live Site Data",      "07_corpus_live_data.md"),
]

EXTRACT_MAP = {
    "Fee Structure": "Fees",
    "Contact":       "Contact Details",
}


def extract_subsection(md_text: str, heading: str) -> str:
    lines = md_text.splitlines()
    collecting = False
    result = []
    pattern = re.compile(r"^##\s+(.+)$")
    for line in lines:
        m = pattern.match(line)
        if m:
            if heading.lower() in m.group(1).lower():
                collecting = True
                result.append(line)
                continue
            elif collecting:
                break
        if collecting:
            result.append(line)
    return "\n".join(result)


def load_kb_markdown_documents() -> list[Document]:
    docs = []
    seen: dict[str, str] = {}
    for section, filename in KB_FILES:
        filepath = os.path.join(KB_DIR, filename)
        if not os.path.exists(filepath):
            print(f"  ⚠️  Missing: {filepath} — skipping '{section}'")
            continue
        if filepath not in seen:
            with open(filepath, encoding="utf-8") as f:
                seen[filepath] = f.read()
        content = seen[filepath]
        extract_heading = EXTRACT_MAP.get(section)
        if extract_heading:
            content = extract_subsection(content, extract_heading)
        if not content.strip():
            print(f"  ⚠️  Empty content for '{section}' — skipping")
            continue
        docs.append(Document(
            page_content=content,
            metadata={"section_hint": section, "source": filename}
        ))
    print(f"  Loaded {len(docs)} section documents from {KB_DIR}")
    return docs


# ── Metadata helpers ───────────────────────────────────────────────────────────

def infer_section(text: str) -> str:
    text_lower = text.lower()
    scores = {}
    keyword_map = {
        "About BVRIT":        ["naac", "aicte", "vision", "mission", "established", "about bvrit",
                               "jntuh", "autonomous", "sri vishnu educational society",
                               "women empowerment", "chairman", "accreditation"],
        "Departments":        ["department", "cse", "ece", "eee", "it", "bs&h", "aiml",
                               "b.tech", "m.tech", "intake", "undergraduate", "postgraduate",
                               "computer science", "electronics"],
        "Admissions":         ["admission", "eamcet", "ecet", "seat", "category a", "category b",
                               "tseamcet", "tsche", "eapcet", "rank", "hostel", "transport"],
        "Fee Structure":      ["fee", "tuition", "₹", "lakh", "per year", "batch", "rupees"],
        "Placements":         ["placement", "lpa", "recruiter", "tap cell", "package",
                               "median salary", "placed", "multinational", "mnc",
                               "campus recruitment"],
        "Campus & Facilities":["library", "gym", "cafeteria", "facility", "research",
                               "innovation", "patent", "infrastructure", "lab", "laboratory",
                               "sports", "hostel capacity", "hostel security", "lady warden",
                               "anti-ragging", "security officer", "500+ students",
                               "4 blocks", "150+ rooms"],
        "Faculty":            ["faculty", "professor", "hod", "principal", "ph.d",
                               "experience", "publications", "dr.", "mr.", "ms.",
                               "associate professor", "assistant professor",
                               "k.v.n. sunitha", "sunitha", "founder principal",
                               "b.tech (ece", "m.tech (cs", "guided", "ph.d scholars",
                               "textbooks", "dst-tide", "dst-nims", "29 years", "30+ years",
                               "patents filed", "research areas"],
        "Contact":            ["contact", "phone", "email", "address", "fax",
                               "hyderabad – 500", "+91 40", "info@bvrit",
                               "bvrithyderabad.edu.in"],
        "Live Site Data":     ["source: https://bvrithyderabad", "live site data",
                               "crawled directly", "post author", "post published",
                               "post category", "flash new"],
    }
    for section, keywords in keyword_map.items():
        scores[section] = sum(1 for kw in keywords if kw in text_lower)

    bio_markers = ["founder principal", "b.tech (ece", "m.tech (cs", "ph.d scholars",
                   "guided", "k.v.n. sunitha", "sunitha", "dst-tide"]
    if scores.get("Faculty", 0) > 0 and any(m in text_lower for m in bio_markers):
        scores["Faculty"] += 3

    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "General"


def infer_subsection(text: str) -> str:
    for line in text.splitlines():
        line = line.strip().lstrip("#").strip()
        if line and len(line) < 80 and line[0].isupper() and not line.endswith("."):
            return line
    return ""


# ── Core build ─────────────────────────────────────────────────────────────────

def build_index():
    # Choose source: docx preferred, markdown fallback
    if os.path.exists(DOCX_PATH):
        print(f"Loading knowledge base from docx: {DOCX_PATH}")
        docs = load_docx_documents()
    else:
        print(f"⚠️  Docx not found — falling back to KB markdown files in {KB_DIR}")
        docs = load_kb_markdown_documents()

    total_chars = sum(len(d.page_content) for d in docs)
    print(f"  Total chars across all sections: {total_chars:,}")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n## ", "\n### ", "\n\n", "\n", ". ", " "],
    )
    chunks = splitter.split_documents(docs)
    print(f"  Split into {len(chunks)} chunks")

    # Attach metadata per chunk
    section_counters: dict[str, int] = {}
    for chunk in chunks:
        hint    = chunk.metadata.get("section_hint", "")
        section = hint if hint else infer_section(chunk.page_content)
        subsection = infer_subsection(chunk.page_content)
        section_counters[section] = section_counters.get(section, 0) + 1
        idx  = section_counters[section]
        slug = section.lower().replace(" & ", "_").replace(" ", "_")
        chunk.metadata.update({
            "source":        chunk.metadata.get("source", "kb"),
            "section":       section,
            "subsection":    subsection,
            "chunk_id":      f"{slug}_{idx:03d}",
            "last_verified": LAST_VERIFIED,
        })

    print("\nChunk distribution by section:")
    for sec, count in sorted(section_counters.items()):
        print(f"  {sec:<30} {count:>3} chunks")

    print("\nEmbedding and persisting to Chroma …")
    embed_kwargs = {"model": EMBED_MODEL}
    if OPENAI_BASE_URL:
        embed_kwargs["openai_api_base"] = OPENAI_BASE_URL
    embeddings = OpenAIEmbeddings(**embed_kwargs)

    # Wipe old index before rebuilding to avoid duplicate chunks
    CHROMA_DIR_USE = CHROMA_DIR
    if os.path.isdir(CHROMA_DIR):
        try:
            shutil.rmtree(CHROMA_DIR)
            print(f"  Wiped existing index at {CHROMA_DIR}")
        except PermissionError:
            alt_dir = CHROMA_DIR + "_new"
            if os.path.isdir(alt_dir):
                shutil.rmtree(alt_dir, ignore_errors=True)
            CHROMA_DIR_USE = alt_dir
            print(f"  ⚠️  Old index locked — writing new index to {alt_dir}")
            print(f"  ⚠️  Stop the app, then rename '{alt_dir}' → 'bvrith_chroma_db'")

    vectorstore = Chroma.from_documents(
        chunks, embeddings, persist_directory=CHROMA_DIR_USE
    )
    count = vectorstore._collection.count()
    print(f"\n✅  Indexed {count} chunks → {CHROMA_DIR_USE}")

    if CHROMA_DIR_USE != CHROMA_DIR:
        print(f"\n⚠️  ACTION REQUIRED:")
        print(f"   1. Stop the Streamlit app (Ctrl+C)")
        print(f"   2. In PowerShell run:")
        print(f"      Remove-Item -Recurse -Force bvrith_chroma_db")
        print(f"      Rename-Item bvrith_chroma_db_new bvrith_chroma_db")
        print(f"   3. Run: streamlit run app.py")
    return count


# ── Check ──────────────────────────────────────────────────────────────────────

def check_index():
    embed_kwargs = {"model": EMBED_MODEL}
    if OPENAI_BASE_URL:
        embed_kwargs["openai_api_base"] = OPENAI_BASE_URL
    embeddings = OpenAIEmbeddings(**embed_kwargs)
    vectorstore = Chroma(persist_directory=CHROMA_DIR, embedding_function=embeddings)
    count = vectorstore._collection.count()
    print(f"Index contains {count} chunks in {CHROMA_DIR}")

    test_queries = [
        "What is the fee for CSE?",
        "Who is the principal?",
        "What is the average placement package?",
    ]
    print("\nSmoke-test retrieval:")
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
    for q in test_queries:
        results = retriever.invoke(q)
        sections_hit = [r.metadata.get("section", "?") for r in results]
        print(f"  Q: {q}")
        print(f"     Sections retrieved: {sections_hit}\n")


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--check", action="store_true", help="Check existing index only")
    args = parser.parse_args()

    if args.check:
        check_index()
    else:
        build_index()
