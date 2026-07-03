"""
build_kb_docx.py
----------------
Merges all BVRIT Hyderabad KB markdown files into a single, properly-structured
BVRITH_Knowledge_Base.docx with Heading 1 / Heading 2 styles that the LangChain
chunker can exploit.

Run:  python build_kb_docx.py
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re, os

# ── Section map ──────────────────────────────────────────────────────────────
SECTIONS = [
    {
        "heading": "1. About BVRIT",
        "file": "kb/00_college_overview.md",
    },
    {
        "heading": "2. Departments",
        "file": "kb/02_departments.md",
    },
    {
        "heading": "3. Admissions",
        "file": "kb/01_admissions.md",
        "subsection_stop": None,      # use entire file except fee-only section
    },
    {
        "heading": "4. Fee Structure",
        "file": "kb/01_admissions.md",
        "extract": "Fees",           # pull only the ## Fees subsection
    },
    {
        "heading": "5. Placements",
        "file": "kb/03_placements.md",
    },
    {
        "heading": "6. Campus & Facilities",
        "file": "kb/04_facilities_research_governance.md",
    },
    {
        "heading": "7. Faculty",
        "file": "kb/05_faculty_directory.md",
    },
    {
        "heading": "8. Contact",
        "file": "kb/04_facilities_research_governance.md",
        "extract": "Contact Details",
    },
    {
        "heading": "9. Live Site Data",
        "file": "kb/07_corpus_live_data.md",
    },
]


def read_md(filepath: str) -> str:
    with open(filepath, encoding="utf-8") as f:
        return f.read()


def extract_subsection(md_text: str, heading: str) -> str:
    """Return content under a specific ## heading until the next ## heading."""
    lines = md_text.splitlines()
    collecting = False
    result = []
    pattern = re.compile(r"^##\s+(.+)$")
    for line in lines:
        m = pattern.match(line)
        if m:
            if heading.lower() in m.group(1).lower():
                collecting = True
                result.append(line)
                continue
            elif collecting:
                break  # stop at the next ## heading
        if collecting:
            result.append(line)
    return "\n".join(result)


def add_md_content_to_doc(doc: Document, md_text: str, top_heading: str):
    """
    Parse markdown text and add paragraphs/headings to the docx.
    - Lines starting with ##  → Heading 2
    - Lines starting with ### → Heading 3
    - Blank lines            → paragraph break
    - Everything else        → Normal paragraph
    - | table rows |         → converted to plain paragraph (keeps data readable)
    """
    lines = md_text.splitlines()
    for line in lines:
        stripped = line.strip()

        # Skip the top-level # title (we already added it as Heading 1)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)

        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)

        elif stripped.startswith("|"):
            # Table row → plain text, strip pipes and excess whitespace
            cells = [c.strip() for c in stripped.split("|") if c.strip()]
            if cells and not all(set(c) <= {"-", " "} for c in cells):
                doc.add_paragraph(" | ".join(cells), style="Normal")

        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")

        elif stripped == "" or stripped == "---":
            continue  # skip horizontal rules and blanks

        else:
            doc.add_paragraph(stripped, style="Normal")


def build_docx():
    doc = Document()

    # Title page paragraph
    title = doc.add_heading("BVRIT HYDERABAD College of Engineering for Women", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Knowledge Base — Last verified: 2026-07-02 | Includes live site corpus data")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for sec in SECTIONS:
        print(f"  Adding section: {sec['heading']}")
        doc.add_heading(sec["heading"], level=1)

        md_text = read_md(sec["file"])

        if "extract" in sec:
            # Only pull the named subsection
            content = extract_subsection(md_text, sec["extract"])
        else:
            content = md_text

        add_md_content_to_doc(doc, content, sec["heading"])
        doc.add_page_break()

    out_path = "BVRITH_Knowledge_Base.docx"
    temp_path = out_path + ".tmp"

    # Save to temp file first to avoid lock issues
    try:
        doc.save(temp_path)
        # Replace the original (works even if original is locked for read)
        import shutil
        shutil.move(temp_path, out_path)
        print(f"\n✅  Saved: {out_path}")
        print(f"   Sections written: {len(SECTIONS)}")
    except Exception as e:
        print(f"\n❌  Failed to save {out_path}: {e}")
        print("   Close any programs that have the file open and try again.")
        if os.path.exists(temp_path):
            os.remove(temp_path)
        raise


if __name__ == "__main__":
    build_docx()
